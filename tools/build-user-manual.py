#!/usr/bin/env python3
"""Build the Korean Samsung Switch Watch v0.11.3 operator manual.

The manual is intentionally generated from sanitized, deterministic WPF
screenshots. It never needs a company switch, a real IP address, or a secret.
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


VERSION = "0.11.3-poc"
DOCUMENT_DATE = "2026-08-05"
FONT = "맑은 고딕"
MONO = "Consolas"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "16324F"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
PALE_GREEN = "ECFDF3"
PALE_GOLD = "FFF8E8"
PALE_RED = "FEF2F2"
GREEN = "167C3A"
GOLD = "8A5A00"
RED = "B42318"
WHITE = "FFFFFF"
TABLE_BORDER = "CBD5E1"

TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS_DXA = {"top": 80, "start": 120, "bottom": 80, "end": 120}


def set_run_font(run, *, name=FONT, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_left_border(paragraph, color, size=16):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_table_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_margins(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in CELL_MARGINS_DXA.items():
        item = tc_mar.find(qn(f"w:{side}"))
        if item is None:
            item = OxmlElement(f"w:{side}")
            tc_mar.append(item)
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    if sum(widths_dxa) != TABLE_WIDTH_DXA:
        raise ValueError(f"Table widths must total {TABLE_WIDTH_DXA} DXA: {widths_dxa}")
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color=TABLE_BORDER, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_image_alt_text(inline_shape, title, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_field(paragraph, instruction, display=""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def _next_numbering_id(numbering):
    existing = [
        int(item.get(qn("w:numId")))
        for item in numbering.findall(qn("w:num"))
        if item.get(qn("w:numId"))
    ]
    return max(existing, default=0) + 1


def _next_abstract_id(numbering):
    existing = [
        int(item.get(qn("w:abstractNumId")))
        for item in numbering.findall(qn("w:abstractNum"))
        if item.get(qn("w:abstractNumId"))
    ]
    return max(existing, default=0) + 1


def create_numbering(doc, kind, *, levels=1):
    numbering = doc.part.numbering_part.element
    abstract_id = _next_abstract_id(numbering)
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel" if levels > 1 else "singleLevel")
    abstract.append(multi)

    for level in range(levels):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)

        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
        lvl.append(num_fmt)

        lvl_text = OxmlElement("w:lvlText")
        if kind == "bullet":
            lvl_text.set(qn("w:val"), "•")
        elif levels > 1:
            lvl_text.set(
                qn("w:val"),
                ".".join(f"%{index + 1}" for index in range(level + 1)) + ".",
            )
        else:
            lvl_text.set(qn("w:val"), "%1.")
        lvl.append(lvl_text)

        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        suffix = OxmlElement("w:suff")
        # Headings read best with a literal space after "1."/"1.1.".
        # Lists keep a tab so wrapped lines align under their text.
        suffix.set(qn("w:val"), "space" if levels > 1 else "tab")
        lvl.append(suffix)

        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        # The numbering suffix tab must land on the paragraph's text indent.
        # A stop before the indent makes a wrapped list number appear on the
        # previous visual line in Microsoft Word.
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)

        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT)
        r_fonts.set(qn("w:hAnsi"), FONT)
        r_fonts.set(qn("w:eastAsia"), FONT)
        r_pr.append(r_fonts)
        lvl.append(r_pr)
        abstract.append(lvl)

    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(list(numbering).index(first_num), abstract)
    num_id = _next_numbering_id(numbering)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level=0):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])


def define_styles(doc):
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(8.5)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = False

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = MONO
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string("243447")
    code_style._element.rPr.rFonts.set(qn("w:ascii"), MONO)
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), MONO)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.08)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing = 1.05


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_header_footer(section, *, first=False):
    configure_page(section)
    section.different_first_page_header_footer = first

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.clear()
    run = p.add_run("Samsung Switch Watch  |  사용자 설명서")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run(f"v{VERSION}  |  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(p, "PAGE", "1")

    if first:
        first_header = section.first_page_header
        first_header.paragraphs[0].clear()
        first_footer = section.first_page_footer
        fp = first_footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        run = fp.add_run(f"v{VERSION}  |  {DOCUMENT_DATE}")
        set_run_font(run, size=8.5, color=MUTED)


def add_body(doc, text, *, bold_prefix=None, center=False):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_callout(
    doc,
    label,
    text,
    kind="info",
    *,
    page_break_before=False,
):
    palette = {
        "info": (LIGHT_BLUE, BLUE),
        "success": (PALE_GREEN, GREEN),
        "warning": (PALE_GOLD, GOLD),
        "danger": (PALE_RED, RED),
    }
    fill, accent = palette[kind]
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8, line=1.18)
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.keep_together = True
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.04)
    set_paragraph_shading(p, fill)
    set_paragraph_left_border(p, accent)
    run = p.add_run(label + "  ")
    set_run_font(run, size=10.5, color=accent, bold=True)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=NAVY)
    return p


def add_code_block(doc, text):
    p = doc.add_paragraph(style="Code Block")
    p.paragraph_format.keep_together = True
    set_paragraph_shading(p, LIGHT_GRAY)
    set_paragraph_left_border(p, MUTED, size=10)
    lines = text.strip("\n").splitlines()
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, name=MONO, size=8.5, color="243447")
    return p


def add_bullets(
    doc,
    items,
    bullet_num_id,
    *,
    font_size=11,
    space_after=6,
    line=1.25,
):
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, bullet_num_id)
        set_paragraph_spacing(p, after=space_after, line=line)
        run = p.add_run(item)
        set_run_font(run, size=font_size)


def add_steps(doc, items):
    num_id = create_numbering(doc, "decimal")
    for item in items:
        p = doc.add_paragraph()
        apply_numbering(p, num_id)
        run = p.add_run(item)
        set_run_font(run)


def add_table(
    doc,
    headers,
    rows,
    widths_dxa,
    *,
    header_size=9,
    body_size=9,
    body_line=1.12,
):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    table.rows[0]._tr.get_or_add_trPr()
    set_repeat_table_header(table.rows[0])
    prevent_table_row_split(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Keep a table heading with at least the first data row instead of
        # leaving an isolated header at the bottom of a page.
        paragraph.paragraph_format.keep_with_next = True
        set_paragraph_spacing(paragraph, after=0, line=1.1)
        run = paragraph.add_run(header)
        set_run_font(run, size=header_size, color=NAVY, bold=True)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_table_row_split(row)
        cells = row.cells
        for index, value in enumerate(values):
            cell = cells[index]
            if row_index % 2:
                set_cell_shading(cell, "FAFBFC")
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if index == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_paragraph_spacing(paragraph, after=0, line=body_line)
            run = paragraph.add_run(str(value))
            set_run_font(run, size=body_size)
        set_table_geometry(table, widths_dxa)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)
    return table


def add_heading(doc, text, level, heading_num_id, *, page_break_before=None):
    p = doc.add_paragraph(style=f"Heading {level}")
    if page_break_before is None:
        page_break_before = level == 1
    p.paragraph_format.page_break_before = page_break_before
    apply_numbering(p, heading_num_id, level - 1)
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_unnumbered_heading(doc, text, level=1, *, page_break_before=False):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break_before
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 16, 2: 13, 3: 12}[level],
        color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level],
        bold=True,
    )
    return p


def add_image(doc, path, *, width, title, alt_text, caption):
    if not path.is_file():
        raise FileNotFoundError(f"Manual screenshot is missing: {path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_image_alt_text(shape, title, alt_text)
    cp = doc.add_paragraph(style="Caption")
    cp.add_run(caption)
    return shape


def add_cover(doc):
    for _ in range(5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("NETWORK OPERATIONS GUIDE")
    set_run_font(run, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("Samsung Switch Watch")
    set_run_font(run, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("원격 삼성 스위치 점검")
    set_run_font(run, size=15, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(30)
    run = p.add_run("Viewer 중심 운영 사용자 설명서")
    set_run_font(run, size=15, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(62)
    run = p.add_run("초급 네트워크 엔지니어부터 상급 관리자까지")
    set_run_font(run, size=10.5, color=GOLD)


def build_manual(output_path: Path, images_dir: Path):
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = False
    section = doc.sections[0]
    configure_header_footer(section, first=True)
    define_styles(doc)

    core = doc.core_properties
    core.title = "Samsung Switch Watch 사용자 설명서"
    core.subject = f"{VERSION} Viewer 중심 원격 삼성 스위치 점검 운영 가이드"
    core.author = "Samsung Switch Watch Project"
    core.keywords = "Samsung Switch Watch, Telnet, Windows Service, Viewer, Agent"
    core.comments = "Sanitized offline manual; contains no company credentials or device data."

    bullet_num_id = create_numbering(doc, "bullet")
    heading_num_id = create_numbering(doc, "decimal", levels=3)

    add_cover(doc)
    add_unnumbered_heading(
        doc,
        "이 설명서에서 가장 먼저 알아둘 것",
        page_break_before=True,
    )
    add_callout(
        doc,
        "핵심 구조",
        "Agent는 스위치에 접속해 결과만 돌려주는 창 없는 Windows 서비스입니다. "
        "장비 IP와 계정, 명령 입력, 감시 주기와 화면 표시는 모두 Viewer가 담당합니다.",
        "info",
    )
    add_table(
        doc,
        ["구성요소", "하는 일", "저장하는 것"],
        [
            ("Agent", "Viewer 요청을 받아 Telnet/23 접속 후 명령 실행", "장비/계정/명령 결과를 저장하지 않음"),
            ("Viewer", "장비 등록, 로그인 확인, show 명령, 주기 감시, 화면 표시", "DPAPI 보호 계정, 기준 해시, 이벤트"),
            ("스위치", "IES4224GP · IES4028XP · IES4226XP", "기존 장비 설정과 로그"),
        ],
        [1600, 3900, 3860],
    )
    add_unnumbered_heading(doc, "3분 빠른 시작", 2)
    add_steps(
        doc,
        [
            "원격 PC에서 Agent ZIP을 풀고 SamsungSwitchWatch.Agent.Setup.exe를 실행한 뒤 UAC를 승인합니다.",
            "Agent Setup에는 Viewer IP나 관리망 CIDR을 입력하지 않습니다. 설치 내용을 확인하고 '설치 / 업데이트'를 누릅니다.",
            "Viewer PC에서 Viewer ZIP을 풀고 SamsungSwitchWatch.Viewer.exe를 직접 실행합니다.",
            "Viewer의 Agent 연결에서 Agent PC 주소만 입력하고 '연결 확인 및 저장'을 누릅니다.",
            "장비 관리에서 장비명, 모델, IPv4, ID, 로그인 PW, 선택 사항인 enable PW를 입력합니다.",
            "로그인 확인이 성공하면 저장하고, 수집 진단에서 포트·로그 명령을 확인한 뒤 필요할 때 주기 감시를 켭니다.",
            "대시보드의 장비 명령 탭에서 한 줄 show 명령을 실행하고 결과를 확인합니다.",
        ],
    )
    add_callout(
        doc,
        "중요",
        "Viewer가 종료되면 주기 감시도 중단됩니다. Agent만 실행 중이라고 감시가 계속되는 구조가 아닙니다.",
        "warning",
    )
    add_heading(doc, "운영 구조 이해", 1, heading_num_id)
    add_body(
        doc,
        "Viewer가 요청할 때마다 Agent는 새 Telnet 세션을 열고 로그인, 필요 시 enable, "
        "명령 실행, 로그아웃 순서로 처리합니다. 한 번의 요청이 끝나면 세션을 닫으므로 "
        "장비의 exec-timeout이 5분이어도 유휴 세션을 계속 붙잡지 않습니다.",
    )
    add_callout(
        doc,
        "명령 시간 제한",
        "포트 상태와 시스템 로그는 순차적인 개별 세션으로 수집합니다. 명령 시간 초과·출력 한도는 "
        "실패한 항목만 확인 불가로 표시하고 다른 항목을 계속 확인합니다. 명령은 30초 동안 새 응답이 "
        "없으면 중단하고, 출력이 계속되어도 전체 90초를 넘기지 않습니다. 같은 실패 명령은 현재 "
        "주기에서 즉시 재시도하지 않습니다. 인증·enable·TCP·세션 종료는 반복 로그인을 막기 위해 "
        "해당 장비 주기를 중단합니다. 이번 POC의 자동 감시 검증·지원 범위는 등록 장비 10대 이하입니다.",
        "info",
    )
    add_code_block(
        doc,
        """
Viewer PC                 Agent PC                    Samsung Switch
장비·계정·명령 입력  ──HTTPS/18443──>  무상태 실행기  ──Telnet/23──>  show 실행
결과·변경 화면 표시  <──────────────  원문 반환       <───────────  출력 반환
""",
    )
    add_table(
        doc,
        ["구간", "고정 통신", "운영 제한"],
        [
            ("Viewer → Agent", "HTTPS/TCP 18443", "10/8, 172.16/12, 192.168/16 사설 주소에서 접근"),
            ("Agent → Switch", "Telnet/TCP 23", "같은 RFC1918 사설 주소의 장비만 허용"),
        ],
        [1900, 2100, 5360],
    )
    add_bullets(
        doc,
        [
            "Agent는 Windows 서비스로 실행되어 RDP 종료나 사용자 로그오프 뒤에도 계속 대기합니다.",
            "Agent는 장비 목록, Telnet 계정, 수동 명령 원문, 스위치 출력 원문을 보관하지 않습니다.",
            "일반 사용자가 볼 수 있는 Agent 창이나 트레이 아이콘은 없습니다.",
            "Viewer는 압축 해제한 폴더에서 포터블로 실행하며, 현재 Windows 사용자 범위에서만 계정을 복호화합니다.",
        ],
        bullet_num_id,
    )
    add_callout(
        doc,
        "보안 전제",
        "Telnet 자체는 암호화되지 않습니다. Agent와 스위치 사이 경로는 반드시 제한된 관리망으로 구성하세요.",
        "danger",
    )
    add_heading(doc, "Agent와 Viewer 설치·연결", 1, heading_num_id)
    add_heading(doc, "Agent 설치", 2, heading_num_id)
    add_callout(
        doc,
        "반입 파일 확인",
        f"{VERSION}는 코드 서명되지 않은 시험판입니다. 압축을 풀기 전에 Agent와 Viewer ZIP의 "
        "SHA-256을 해당 GitHub Release 본문에 표시된 값과 비교하고 사내 보안 반입 절차를 "
        "완료하세요. SmartScreen 또는 EDR 경고를 우회하지 않습니다.",
        "danger",
    )
    add_steps(
        doc,
        [
            "Agent 릴리스 ZIP을 원격 PC의 임시 폴더에 압축 해제합니다.",
            "SamsungSwitchWatch.Agent.Setup.exe를 실행하고 UAC 관리자 승인을 합니다.",
            "설치 내용을 확인합니다. Viewer와 스위치 허용 범위는 RFC1918 사설 대역으로 자동 적용되므로 별도 입력이 없습니다.",
            "Setup이 중단된 이전 설치 기록을 발견하면 설치/업데이트가 잠깁니다. '이전 상태 복구'를 먼저 누르고 새 작업 기록 검사까지 완료됐는지 확인합니다.",
            "'설치 / 업데이트'를 누르면 자동 점검과 설치가 이어집니다. 완료 또는 '설치 완료 · 연결 확인 필요' 경고를 확인합니다.",
        ],
    )
    add_image(
        doc,
        images_dir / "00-agent-setup.png",
        width=4.5,
        title="Agent Setup 화면",
        alt_text="Viewer IP와 관리망 CIDR 입력 없이 설치 내용과 설치 버튼을 보여 주는 Agent Setup 화면",
        caption="그림 1. 별도 네트워크 입력이 없는 Agent Setup 화면",
    )
    add_callout(
        doc,
        "입력 없이 설치",
        "Setup은 제품 파일, Windows 서비스와 방화벽 규칙을 준비합니다. Viewer IP와 스위치 CIDR은 "
        "요구하지 않습니다. Agent는 10/8, 172.16/12, 192.168/16에서 들어오는 Viewer 요청과 "
        "같은 사설 대역의 Telnet/23 장비만 자동으로 허용합니다.",
        "info",
    )
    add_image(
        doc,
        images_dir / "00-agent-setup-recovery-failed.png",
        width=4.5,
        title="Agent Setup 복구 실패 화면",
        alt_text="복구 자료 정리가 끝나지 않아 설치를 계속 잠그고 작업 기록 보존, 복구 재시도와 익명 진단 저장을 안내하는 Agent Setup 화면",
        caption="그림 2. 복구 자료 정리 실패와 설치 차단 화면",
    )
    add_callout(
        doc,
        "복구 자료 정리가 끝나지 않은 경우",
        "화면 상단에 SETUP_ROLLBACK_FAILED가 표시되고 복구 자료 정리 분류 아래 "
        "ROLLBACK_JOURNAL_CLEANUP_FAILED 같은 대상별 코드가 표시되면, staging·backup·failed·journal 중 "
        "검증된 정리 대상을 최대 3회 시도하고 실패한 시도 사이 250ms 대기한 뒤에도 남아 있다는 뜻입니다. "
        "설치 버튼은 계속 잠기며 화면에서 '복구 다시 시도'를 사용할 수 있습니다. 반복되면 "
        "'익명 진단 저장'을 사용하고, 폴더나 journal을 수동으로 삭제하지 마세요.",
        "danger",
    )
    add_callout(
        doc,
        "짧은 SWD1 지원 코드",
        "실패하면 '지원 코드 · 이 코드만 전달하세요' 아래 "
        "SWD1-XXXX-XXXX-XXXX-XXXX 형식의 읽기 전용 코드가 표시됩니다. 별도 복사 버튼은 "
        "없으며 코드만 선택해 Ctrl+C로 복사합니다. 전화나 메신저로 짧은 장애 분류를 전달할 "
        "때 사용하세요. 이 코드는 오프라인에서 생성되고 CRC로 입력 오타를 확인하지만 비밀값, "
        "인증·페어링 토큰 또는 인증서 지문이 아닙니다. 로컬 HTTPS의 다섯 세부 원인은 기존 "
        "SWD1/1 형식과 호환되도록 HTTPS 요청 실패라는 상위 분류로만 담깁니다. 정확한 세부 "
        "코드와 안전 관측값은 같은 실패 화면의 진단정보를 함께 확인하세요. 새 작업과 성공 "
        "화면에서는 이전 코드가 지워집니다.",
        "warning",
    )
    add_callout(
        doc,
        "기존 TCP/18443 허용 규칙 경고",
        "다른 프로그램이 만든 인바운드 허용 규칙이 있으면 "
        "FIREWALL_OVERLAP_PROTECTED 경고를 표시하지만 해당 규칙은 변경하지 않고 설치를 "
        "계속합니다. 제품 규칙은 RFC1918 사설 원격 주소, Domain/Private 프로필과 TCP/18443을 "
        "대상으로 시도합니다.",
        "warning",
    )
    add_callout(
        doc,
        "방화벽 확인 경고는 Agent 설치 실패가 아님",
        "제품 규칙의 적용·재조회 또는 회사 GPO 확인이 실패해도 Agent 서비스와 프로그램은 "
        "설치된 상태로 유지됩니다. 화면에는 설치 실패가 아니라 연결 확인 경고가 표시됩니다. "
        "Viewer에서 연결 테스트를 실행하고 TCP/18443이 실패할 때만 Windows 관리자에게 "
        "방화벽·GPO 경로 확인을 요청하세요.",
        "warning",
    )
    add_callout(
        doc,
        "방화벽은 보조 수단",
        "방화벽 규칙은 원격 연결을 돕기 위한 최선 노력 단계입니다. 규칙 생성 실패가 프로그램 "
        "파일과 서비스 설치를 취소하지 않습니다. Agent 자체도 Viewer와 장비 주소가 RFC1918 "
        "사설 범위인지 다시 확인합니다.",
        "info",
    )
    add_code_block(
        doc,
        """
서비스 이름          : SamsungSwitchWatchAgent
서비스 계정          : NT SERVICE\\SamsungSwitchWatchAgent
통신 포트            : HTTPS/TCP 18443
Viewer 허용 범위      : 10/8, 172.16/12, 192.168/16
장비 허용 범위        : 10/8, 172.16/12, 192.168/16 · Telnet/23
        """,
    )
    add_callout(
        doc,
        "고정 설치 위치",
        "Agent 프로그램은 %ProgramFiles%\\SamsungSwitchWatch\\Agent에 둡니다. 실행할 때마다 "
        "새 HTTPS 인증서를 만들고 Windows Schannel용 임시 사용자 키 컨테이너를 프로세스 수명 동안만 "
        "사용합니다. Agent 종료 시 이를 정리하며 영구 인증서 지문이나 페어링 자료는 저장하지 않습니다. "
        "Setup이 경로 신뢰 검사를 통과하지 못하면 "
        "설치를 중단합니다. 폴더를 강제로 삭제하거나 소유권을 바꿔 우회하지 말고 Windows "
        "관리자에게 실패 코드를 전달하세요.",
        "danger",
    )
    add_unnumbered_heading(
        doc,
        "설치 안정성과 진단",
        level=2,
        page_break_before=True,
    )
    add_callout(
        doc,
        "업데이트 안정성",
        "Setup은 BUILD-MANIFEST.json과 Agent 실행 파일 SHA-256을 확인하고, Program Files의 "
        "임시 staging에 복사한 파일을 다시 검사한 뒤 교체합니다. 패키지·파일·서비스 변경 자체가 "
        "실패한 경우에만 설치 전 상태로 되돌리기를 시도합니다. 서비스가 설치·시작된 뒤 로컬 "
        "HTTPS/API/버전 또는 방화벽 준비 상태를 확인하지 못한 경우에는 설치를 되돌리지 않고 "
        "AGENT_LOCAL_CONNECTION_UNCONFIRMED 경고로 남깁니다. "
        "다음 실행에서 완료되지 않은 "
        "journal이 남아 있으면 자동 복구하지 않고 별도의 '이전 상태 복구' 작업을 요구합니다. "
        "검사 중에는 현재 서비스 "
        "PID와 TCP/18443 소유를 매번 다시 확인합니다. rollback 전에는 관찰한 서비스 프로세스 "
        "종료를 확인하고 프로그램 폴더의 일시적 잠금만 최대 5회 제한적으로 재시도합니다. "
        "staging·backup·failed·journal 정리가 잠시 실패하면 정확히 검증된 대상만 최대 3회 "
        "시도하고 실패한 시도 사이 250ms 대기한 뒤 삭제 결과를 확인합니다. "
        "변경 자체의 복구까지 실패하면 최초 설치 실패 코드를 별도로 유지하고 SETUP_ROLLBACK_FAILED와 "
        "실패한 ROLLBACK_* 단계를 함께 표시합니다.",
        "info",
    )
    add_callout(
        doc,
        "서비스 조회는 한 번만 짧게 재시도",
        "설치/업데이트는 별도 사전 점검과 배포에서 같은 서비스 정보를 중복 조회하지 않습니다. "
        "초기 조회가 일시적으로 실패하면 200ms 뒤 한 번만 다시 확인합니다. 계속 실패하면 "
        "파일·서비스·방화벽 변경 전에 SETUP_SERVICE_FAILED로 중단합니다. 기존 서비스의 보안 "
        "설명자만 읽을 수 없는 경우에는 조회 가능한 구성과 상태로 설치를 계속하되 기존 DACL은 "
        "변경하지 않습니다. 설명자를 확보한 경우에만 새 제한 DACL을 적용하고 실패 시 복원합니다.",
        "info",
    )
    add_callout(
        doc,
        "실패할 때만 진단정보 복사",
        "설치 또는 복구가 실패한 경우에만 '진단정보 복사' 버튼이 나타납니다. 복사 내용에는 "
        "프로그램 버전, UTC 시각, 작업 종류, 최초 실패 코드, ROLLBACK_* 코드, journal 단계와 "
        "파일·서비스 존재 여부, 안전한 readiness 하위 원인만 포함됩니다. 실제 서비스 PID, "
        "IP/CIDR, PC·사용자 이름, 절대 경로, 계정, "
        "비밀번호, 인증서, 방화벽 규칙 원문과 장비 명령·출력은 포함하지 않으며 별도 로그 파일로 "
        "저장하지 않습니다.",
        "info",
    )
    add_callout(
        doc,
        "로컬 HTTPS 실패는 Agent PC 내부 구간",
        "HTTPS_TLS_FAILED, HTTPS_REQUEST_TIMEOUT, HTTPS_CONNECTION_RESET, HTTPS_EOF, "
        "HTTPS_CONNECT_FAILED는 Setup → 127.0.0.1:18443 → Agent 서비스 사이의 로컬 "
        "준비 상태 확인 실패입니다. 이 다섯 코드는 Viewer IP나 스위치 관리망 문제를 뜻하지 "
        "않으며 설치된 서비스와 파일을 되돌리는 사유도 아닙니다. Viewer 연결 테스트로 실제 "
        "사용 가능 여부를 확인하고 실패하면 Agent PC의 Windows 이벤트, TLS 정책과 백신·EDR "
        "개입 여부를 점검하세요.",
        "warning",
    )
    add_table(
        doc,
        ["안전 진단 항목", "의미"],
        [
            ("AgentHealthCode", "로컬 HTTPS 실패의 다섯 세부 분류 중 하나"),
            ("ServiceRunningObserved", "검사 중 Agent 서비스 Running을 관찰했는지 여부"),
            ("ListenerOwnedObserved", "기대 서비스가 TCP/18443 listener를 소유함을 관찰했는지 여부"),
            ("HttpAttemptCount", "로컬 HTTPS 요청을 시도한 제한된 횟수"),
            ("LastTransportPhase", "ListenerOwned, RequestStarted, ResponseHeaders 등 마지막 안전 단계"),
            ("AgentRestartObserved", "검사 중 Agent 서비스 프로세스 재시작을 관찰했는지 여부"),
        ],
        [2900, 6460],
        header_size=8.5,
        body_size=8.25,
        body_line=1.0,
    )
    add_code_block(
        doc,
        """
화면 상태 : 설치 완료 · 연결 확인 필요
경고 코드 : AGENT_LOCAL_CONNECTION_UNCONFIRMED
설치 결과 : Agent 서비스와 프로그램은 유지
다음 조치 : Viewer에서 Agent 연결 확인 및 저장
        """,
    )
    add_callout(
        doc,
        "진단만 세분화됨",
        "설치 완료 경고는 파일·서비스 설치 완료와 연결 준비 미확인을 구분합니다. 이 경우 "
        "rollback은 실행하지 않습니다. 익명 진단의 Health 값과 단계 시간은 원인 분류용이며 "
        "현재 상태를 영구 보장하지 않습니다.",
        "info",
    )
    add_callout(
        doc,
        "성공·실패 후 익명 진단 저장",
        "검사, 설치 또는 복구가 끝나면 '익명 진단 저장'으로 SSW_FIELD_DIAGNOSTIC/2 UTF-8 BOM "
        "TXT를 수동 저장할 수 있습니다. 사진 한 장으로 전달할 수 있도록 최대 12줄, 줄당 88자로 "
        "제한하며 제품·Windows 버전, 안전한 단계·결과·오류·조치 코드와 핵심 상태를 보존합니다. "
        "실제 IP/CIDR, PC·사용자명, 계정, 인증서 정보, 절대 경로, 방화벽·예외 원문, 명령과 장비 "
        "출력은 포함하지 않습니다. 과거 /1 파일은 재현 도구에서 계속 분석할 수 있으며 저장 실패는 "
        "DIAGNOSTIC_WRITE_FAILED로 표시됩니다.",
        "info",
    )
    add_callout(
        doc,
        "네트워크 정책은 자동 적용",
        "Viewer PC 주소나 스위치 관리망이 바뀌어도 CIDR을 다시 입력하지 않습니다. Viewer와 "
        "스위치가 10/8, 172.16/12 또는 192.168/16에 있고 Agent까지 라우팅 가능한지 확인합니다. "
        "공인 주소와 그 밖의 특수 주소는 허용되지 않습니다.",
        "info",
    )
    add_unnumbered_heading(
        doc,
        "설치·연결 오류 확인",
        level=2,
        page_break_before=False,
    )
    add_callout(
        doc,
        "설치 경고와 Viewer 연결 거부",
        "Agent Setup에 AGENT_LOCAL_CONNECTION_UNCONFIRMED 또는 FIREWALL_REMOTE_ACCESS_UNCONFIRMED가 "
        "표시되면 Agent 서비스와 프로그램은 설치됐지만 연결 준비를 확인하지 못했다는 뜻입니다. Viewer에서 "
        "연결 테스트를 실행하고, TCP/18443이 실패할 때 Windows 관리자에게 확인하세요. 규칙을 "
        "AGENT_CONNECTION_REFUSED가 보이면 Viewer에 스위치 IP나 localhost가 아니라 "
        "Agent를 설치한 PC의 실제 주소를 입력했는지 먼저 확인하세요. 같은 PC 시험에서는 "
        "해당 PC의 사설 IPv4 또는 localhost를 사용할 수 있습니다. 그 다음 "
        "Windows 서비스에서 SamsungSwitchWatchAgent가 Running인지 확인하고 TCP/18443을 점검합니다. "
        "TCP 단계는 성공하고 HTTPS 단계가 실패하면 방화벽이 아니라 Agent PC의 로컬 HTTPS/TLS "
        "준비 상태를 확인합니다. "
        "AGENT_CLIENT_NOT_ALLOWED이면 Viewer 주소가 RFC1918 사설 범위인지 확인합니다. "
        f"서비스를 수동 등록하지 말고 {VERSION} Agent Setup을 사용하세요.",
        "danger",
    )
    add_heading(doc, "Viewer 실행", 2, heading_num_id)
    add_callout(
        doc,
        "0.9 설치형 Viewer에서 전환",
        "이전 Viewer가 실행 중이면 새 Viewer는 같은 사용자 데이터를 동시에 쓰지 않도록 "
        "실행을 차단하고 전환 순서를 표시합니다. 기존 트레이 아이콘에서 '프로그램 종료'를 "
        "선택한 뒤 Win+R에서 shell:startup을 열어 'Samsung Switch Watch' 바로 가기를 "
        "삭제합니다. 이어 shell:programs에서도 같은 이름의 이전 시작 메뉴 바로 가기를 "
        "삭제하고 새 Viewer를 다시 실행하세요. 기존 Agent 연결과 장비 정보는 유지됩니다. "
        f"창이 바로 보이지 않으면 알림 영역에서 대시보드를 열고 실행 경로가 새 {VERSION} "
        "폴더인지 확인하세요. API v4가 호환되면 Agent와 Viewer의 세부 버전이 달라도 경고만 "
        "표시하고 연결됩니다.",
        "warning",
    )
    add_steps(
        doc,
        [
            "Viewer 릴리스 ZIP을 운영자 PC에서 계속 사용할 로컬 폴더에 압축 해제합니다.",
            "SamsungSwitchWatch.Viewer.exe를 더블클릭합니다.",
            "Viewer는 일반 사용자 권한으로 실행하며 설치, UAC와 Windows 로그인 자동 시작을 수행하지 않습니다.",
        ],
    )
    add_callout(
        doc,
        "프로그램과 사용자 데이터 분리",
        "Viewer는 압축 해제한 폴더에서 실행합니다. 장비 목록, DPAPI 자격 증명, 감시 이력과 화면 설정은 "
        "%LOCALAPPDATA%\\SamsungSwitchWatch에 보존됩니다. Viewer 폴더를 교체해도 자료는 유지되지만 "
        "문제 재현과 지원을 단순하게 하려면 Agent와 같은 Release 사용을 권장합니다. API v4가 "
        "호환되면 세부 버전이 달라도 연결은 유지되고 경고만 표시됩니다.",
        "info",
    )
    add_callout(
        doc,
        "Viewer가 실행되지 않는 경우",
        f"{VERSION}는 코드 서명되지 않아 SmartScreen이나 사내 EDR이 차단할 수 있습니다. ZIP을 완전히 "
        "압축 해제하고 Viewer EXE와 제공된 DLL이 같은 폴더에 있는지 확인하세요. EDR·AppLocker·WDAC가 "
        "차단하면 우회하지 말고 공식 ZIP의 버전·해시와 차단 기록을 보안 담당자에게 전달하세요.",
        "danger",
    )
    add_callout(
        doc,
        "포터블 운영",
        "공개 Viewer 패키지에는 설치 스크립트와 자동 시작 기능이 없습니다. 필요하면 사용자가 "
        "직접 바로 가기를 만들고, 감시가 필요할 때 Viewer를 실행해 두세요.",
        "info",
    )
    add_heading(doc, "Viewer에서 Agent 연결", 2, heading_num_id)
    add_image(
        doc,
        images_dir / "02-agent-connection.png",
        width=3.5,
        title="Agent 연결 창",
        alt_text="Agent PC 주소 하나와 자동 HTTPS 연결 단계, 연결 확인 및 저장 버튼을 보여 주는 연결 설정 창",
        caption="그림 3. Agent 주소만 입력하는 연결 설정",
    )
    add_unnumbered_heading(
        doc,
        "연결 실패 화면과 확인 순서",
        level=3,
        page_break_before=True,
    )
    add_image(
        doc,
        images_dir / "02-agent-connection-failed.png",
        width=3.5,
        title="Agent 연결 실패와 지원 코드",
        alt_text="TCP 18443 연결 거부 단계와 선택 가능한 SWD1 지원 코드를 함께 표시하는 Viewer Agent 연결 실패 창",
        caption="그림 4. Viewer Agent 연결 실패와 짧은 SWD1 지원 코드",
    )
    add_bullets(
        doc,
        [
            "Agent를 설치한 원격 PC의 IPv4 또는 사내 DNS 이름만 입력합니다. 스위치 IP나 Viewer PC 주소를 "
            "입력하지 않습니다.",
            "Agent와 Viewer가 같은 PC라면 localhost 또는 해당 PC의 사설 IPv4를 사용할 수 있습니다.",
            "https://, 포트, 인증서 지문과 페어링 토큰은 입력하지 않습니다. HTTPS/TCP 18443과 "
            "Agent 인증서 확인은 프로그램이 자동 처리합니다.",
        ],
        bullet_num_id,
    )
    add_callout(
        doc,
        "동일 PC에서 먼저 확인",
        "Agent와 Viewer를 같은 PC에 설치했다면 Agent PC 주소에 localhost 또는 해당 PC의 사설 "
        "IPv4를 입력하고 일반 연결 확인을 실행합니다. 성공은 Agent 서비스, TCP/18443, HTTPS와 "
        "Agent API만 뜻합니다. 스위치 자격 증명과 명령은 사용하지 않으며, 장비는 저장 후 "
        "'장비 관리 → 로그인 확인'에서 계정과 프롬프트를 확인하고 수집 진단에서 실제 명령을 "
        "확인합니다. 실제 원격 배치 전에는 원격 Viewer PC에서 "
        "Agent PC 주소로 연결 확인을 다시 실행하세요.",
        "info",
    )
    add_unnumbered_heading(
        doc,
        "연결 진단과 지원 정보",
        level=2,
        page_break_before=False,
    )
    add_callout(
        doc,
        "Agent 연결 익명 진단",
        "Agent 연결 확인이 성공 또는 실패로 끝나면 '익명 진단 저장'을 사용할 수 "
        "있습니다. 사진 한 장용 최대 12줄 TXT에는 주소·DNS·TCP/18443·HTTPS·Agent API 단계 상태와 "
        "제한된 소요 시간과 확인된 Agent/API 버전만 기록하며 입력 주소, DNS 이름, "
        "PC·사용자명, 계정, 경로, 예외 원문과 장비 명령·출력은 제외합니다. 연결 성공 뒤 창은 "
        "저장 결과를 보여 주며, 필요하면 진단을 저장한 다음 닫습니다.",
        "info",
    )
    add_callout(
        doc,
        "세 가지 진단을 구분하세요",
        "SWD1은 Agent Setup 또는 Viewer 연결 실패를 전화·메신저로 짧게 전달하는 코드입니다. "
        "Agent Setup의 '진단정보 복사'는 실패 전용 긴 비식별 요약을 클립보드에 복사하고, "
        "'익명 진단 저장'은 작업이나 연결 검사가 끝난 뒤 SSW_FIELD_DIAGNOSTIC/2 한 장용 TXT를 "
        "사용자가 선택해 저장합니다. SWD1은 기존 진단을 대체하지 않으며 접속 권한을 부여하지 "
        "않습니다.",
        "warning",
    )
    add_heading(
        doc,
        "장비와 계정 등록",
        1,
        heading_num_id,
        page_break_before=False,
    )
    add_image(
        doc,
        images_dir / "03-device-management.png",
        width=5.5,
        title="장비 관리 창",
        alt_text="장비명, 모델, IPv4, 계정 ID, 로그인 비밀번호, enable 비밀번호와 감시 설정을 입력하는 창",
        caption="그림 5. Viewer가 보관하는 장비 및 계정 입력 화면",
    )
    add_unnumbered_heading(
        doc,
        "장비 입력 항목",
        level=3,
        page_break_before=True,
    )
    add_table(
        doc,
        ["입력 항목", "필수", "설명"],
        [
            ("장비명", "예", "운영자가 구분하기 쉬운 표시 이름"),
            ("모델", "예", "IES4224GP, IES4028XP, IES4226XP"),
            ("장비 IPv4", "예", "10/8, 172.16/12 또는 192.168/16의 스위치 관리 IPv4"),
            ("계정 ID", "예", "Telnet 로그인 계정"),
            ("로그인 PW", "예", "현재 Windows 사용자 DPAPI로 보호"),
            ("enable PW", "아니요", "로그인 후 프롬프트가 >인 장비에서만 사용"),
            ("주기 감시", "아니요", "기본값 꺼짐, 로그인 확인 성공 후 켤 수 있음"),
        ],
        [1900, 900, 6560],
    )
    add_callout(
        doc,
        "로그인 확인 실패",
        "실패한 장비도 저장할 수 있지만 '로그인 미확인'으로 표시되고 주기 감시는 강제로 꺼집니다. "
        "주소가 사설 대역인지, ID/PW와 enable 필요 여부를 바로잡은 뒤 다시 확인하세요. 성공해도 "
        "조회 명령까지 검증된 것은 아니므로 수집 진단을 이어서 실행합니다.",
        "warning",
    )
    add_heading(doc, "장비 show 명령 실행", 1, heading_num_id)
    add_body(
        doc,
        "대시보드에서 장비를 선택하고 '장비 명령' 탭에 한 줄짜리 show 명령을 입력합니다. "
        "Viewer는 자유로운 조회를 허용하되 설정 모드로 들어갈 수 있는 입력은 보내지 않습니다.",
    )
    add_image(
        doc,
        images_dir / "04-command-output.png",
        width=4.35,
        title="장비 명령 실행 화면",
        alt_text="show port status를 입력하고 데모 스위치의 익명화된 결과를 확인하는 장비 명령 탭",
        caption="그림 6. 한 줄 show 명령 실행과 메모리 내 결과 확인",
    )
    add_table(
        doc,
        ["구분", "예시", "처리"],
        [
            ("허용", "한 줄짜리 show ...", "조회 명령 실행"),
            ("허용", "show port status", "실행, 포트 상태 확인"),
            ("차단", "구분자 또는 설정 명령", "한 줄 show 명령이 아님"),
        ],
        [1300, 3500, 4560],
    )
    add_bullets(
        doc,
        [
            "Enter: 실행, Esc: 취소. 출력 상한은 64KiB이며 복사 버튼은 Windows 클립보드만 사용합니다.",
            "완료 줄에는 처리 시간과 세션 횟수, 재연결이 있었다면 재연결 횟수도 표시됩니다.",
            "수동 명령 원문과 출력은 Viewer 프로세스가 종료되면 사라지며 파일이나 Agent에 저장되지 않습니다.",
        ],
        bullet_num_id,
    )
    add_callout(
        doc,
        "민감한 show 명령",
        "결과에 비밀번호 해시, SNMP 문자열, IP/VLAN과 망 구성이 포함될 수 있습니다. "
        "복사한 출력은 사내 보안 기준에 따라 취급하고 메신저나 일반 문서로 전달하지 마세요.",
        "danger",
    )
    add_heading(doc, "Viewer 주기 감시", 1, heading_num_id)
    add_body(
        doc,
        "주기 감시는 장비별로 켭니다. Viewer가 실행 중인 동안만 정해진 조회 명령을 요청하고, "
        "이전 기준과 비교해 포트 상태 변화와 새 로그를 이벤트로 만듭니다.",
    )
    add_table(
        doc,
        ["용도", "우선 명령", "대체 처리"],
        [
            ("포트 상태", "show port status", "모델에서 미지원이면 해당 점검만 실패 표시"),
            ("최근 로그", "show sylog tail num 100", "미지원이면 show syslog tail num 100, show log ram 순서로 재시도"),
        ],
        [1800, 3300, 4260],
    )
    add_bullets(
        doc,
        [
            "최초 정상 결과는 기준선으로만 저장하며 기존 로그를 새 이벤트로 알리지 않습니다.",
            "포트 상태는 의미 있는 필드로 비교합니다. 최초 Down은 기존 상태로 보고, 이후 Up → Down과 Down → Up만 표시합니다.",
            "중요 업링크로 지정되지 않은 일반 포트 변화는 장애가 아닌 경고로 표시하며, 케이블·상대 장비와 실제 사용 여부를 확인합니다.",
            "로그는 저장된 식별 해시와 비교해 새 항목만 이벤트로 만듭니다.",
            "주기 감시 저장소에는 기준 해시, 이벤트와 상태만 남고 Telnet 원문은 남지 않습니다.",
            "같은 장애가 계속되면 반복 팝업 대신 지속 상태를 유지하고, 정상화되면 복구 이벤트를 만듭니다.",
        ],
        bullet_num_id,
    )
    add_callout(
        doc,
        "감시 공백",
        "Viewer 종료, PC 절전 또는 네트워크 단절 중에는 감시하지 않습니다. 다시 실행하면 중단 시간을 "
        "감시 공백으로 표시하고 현재 결과를 새 기준으로 사용합니다. 그 사이 로그 버퍼에서 사라진 사건은 복원할 수 없습니다.",
        "warning",
    )
    add_callout(
        doc,
        "장비 목록 파일 오류",
        "장비 목록 파일이 손상됐거나 현재 Viewer보다 새로운 형식이거나 읽기·쓰기가 불가능하면 "
        "Viewer는 원본을 보존하고 해당 실행의 자동 감시와 장비 명령을 중지합니다. 기존 장비가 "
        "정상으로 보이거나 빈 목록으로 덮어써지지 않습니다. 화면의 오류 코드에 따라 격리 파일, "
        "Viewer 버전, 사용자 폴더 권한·파일 잠금과 Windows 사용자 프로필을 확인한 뒤 Viewer를 "
        "다시 시작하세요.",
        "danger",
    )
    add_callout(
        doc,
        "감시 이력 파일 오류",
        "감시 이력 파일이 손상됐거나 현재 Viewer보다 새로운 형식이거나 읽기·쓰기가 불가능하면 "
        "Viewer는 이전 장애를 정상으로 오판하지 않도록 자동 감시와 이력 저장을 중지합니다. "
        "장비의 수동 show 조회는 계속 사용할 수 있습니다. 화면의 오류 코드와 조치 안내에 따라 "
        "격리 파일, Viewer 버전, 사용자 폴더 권한·파일 잠금을 확인한 뒤 Viewer를 다시 시작하세요.",
        "warning",
    )
    add_heading(doc, "짧은 세션 유지 시간 대응", 2, heading_num_id)
    add_code_block(
        doc,
        """
매 요청: TCP 연결 → 로그인 → 필요 시 enable → 명령 실행 → exit/logout → 소켓 종료
최대 세션 시간: 240초
명령 합산 예산이 240초를 넘는 경우: 순서를 유지한 여러 세션으로 미리 분할
명령 단계에서 세션 종료: 완료 명령은 반복하지 않고 남은 명령만 새 세션으로 1회 재시도
자동 재시도 없음: 인증/enable 실패, 명령 타임아웃
완료 표시: 세션 횟수와 재연결 횟수
""",
    )
    add_heading(doc, "대시보드 읽는 법", 1, heading_num_id)
    add_image(
        doc,
        images_dir / "01-dashboard.png",
        width=5.2,
        title="Viewer 대시보드",
        alt_text="장비 목록, 선택 장비 상태, 최근 이벤트와 Viewer 감시 상태를 보여 주는 대시보드",
        caption="그림 7. Viewer 중심 대시보드 전체 화면",
    )
    add_table(
        doc,
        ["영역", "확인할 내용"],
        [
            ("상단", "Agent 연결, 장비 관리, 수동 새로고침, 미니 창"),
            ("요약", "현재 확인/마지막 확인, 정상/경고/장애/연결 끊김/미감시 수"),
            ("왼쪽", "문제 우선으로 정렬된 등록 장비와 마지막 확인 시각"),
            ("가운데", "선택 장비 상태, 새 로그, 변경 이력, 장비 명령"),
            ("오른쪽", "모든 장비의 최근 이벤트, 검색과 확인 처리"),
            ("하단", "현재 작업과 계정·원문 저장 정책"),
        ],
        [1800, 7560],
    )
    add_bullets(
        doc,
        [
            "Agent 연결이 끊겼다고 모든 스위치를 Down으로 바꾸지 않습니다. 마지막 상태를 유지하고 '미확인'으로 표시합니다.",
            "이벤트 확인은 복구가 아니며, CSV/JSON 내보내기는 이벤트만 익명화하고 수동 명령 출력은 제외합니다.",
        ],
        bullet_num_id,
        font_size=10.5,
        space_after=4,
        line=1.18,
    )
    add_heading(doc, "미니 창과 장애 알림", 1, heading_num_id)
    add_body(
        doc,
        "평소에는 미니 창을 항상 위에 두고 문제 수와 마지막 점검 시각만 확인할 수 있습니다. "
        "장애가 새로 발생하면 별도의 팝업이 나타나며 클릭하면 해당 이벤트로 이동합니다.",
    )
    add_image(
        doc,
        images_dir / "05-mini-window.png",
        width=3.25,
        title="항상 위 미니 창",
        alt_text="정상, 경고, 장애 수와 최근 문제를 보여 주는 작은 항상 위 창",
        caption="그림 8. 반복 운영용 미니 창",
    )
    add_image(
        doc,
        images_dir / "06-alert-popup.png",
        width=3.45,
        title="장애 알림 팝업",
        alt_text="데모 업링크 포트 Down 장애와 발생 시각을 보여 주는 알림 팝업",
        caption="그림 9. 새 장애 알림 팝업",
    )
    add_bullets(
        doc,
        [
            "핀 버튼: 다른 일반 창보다 위에 표시",
            "대시보드 열기: 전체 상태와 변경 이력 확인",
            "새로고침: Viewer에서 즉시 상태 요청",
            "같은 장애 지속 중에는 팝업을 반복하지 않고 복구 시 별도 알림",
            "팝업 위에 마우스를 두거나 키보드 포커스가 있으면 읽는 동안 자동으로 닫히지 않음",
        ],
        bullet_num_id,
    )
    add_heading(doc, "저장 위치와 보안 경계", 1, heading_num_id)
    add_table(
        doc,
        ["데이터", "위치", "보호/수명"],
        [
            ("Viewer 프로그램", "사용자가 압축 해제한\n로컬 폴더", "포터블, 일반 사용자 권한으로 실행"),
            ("장비명·모델·IPv4", "Viewer 사용자 프로필", "현재 Windows 사용자 범위"),
            ("ID·로그인 PW·enable PW", "Viewer 사용자 프로필", "DPAPI CurrentUser 암호화"),
            ("주기 감시 기준 해시·이벤트", "Viewer 사용자 프로필", "원문 없이 로컬 저장"),
            ("수동 show 입력·출력", "Viewer 프로세스 메모리", "복사 가능, 종료 시 소멸"),
            (
                "Agent HTTPS 인증서",
                "Agent 프로세스 메모리",
                "서비스 시작마다 새로 생성, 종료 시 소멸",
            ),
            (
                "Agent 실행 설정",
                "%ProgramFiles%\\SamsungSwitchWatch\\Agent",
                "RFC1918 Viewer/장비 정책과 TCP 포트",
            ),
            ("장비 계정·스위치 출력", "Agent", "저장하지 않음"),
        ],
        [2200, 3000, 4160],
    )
    add_bullets(
        doc,
        [
            "Viewer 설정을 다른 PC나 Windows 사용자에게 복사해도 계정은 복호화되지 않으며, 진단 파일에는 "
            "IP·ID·비밀번호·호스트명·수동 명령 원문을 넣지 않습니다.",
            "Viewer는 Agent 인증서 지문이나 페어링 토큰을 저장하지 않습니다. HTTPS는 전송 내용을 "
            "암호화하지만 Agent PC 신원을 별도로 인증하는 구조는 아닙니다.",
            "제품 방화벽 규칙과 Agent 업무 API는 RFC1918 사설 Viewer 주소를 허용합니다. "
            "Agent는 RFC1918 사설 장비의 Telnet/23만 사용합니다.",
        ],
        bullet_num_id,
    )
    add_heading(
        doc,
        "문제 해결",
        1,
        heading_num_id,
        page_break_before=False,
    )
    add_callout(
        doc,
        "운영 원칙",
        "Viewer PC와 Agent PC는 관리망에서만 사용하고, 일반 사용자 VLAN이나 인터넷을 거쳐 Telnet을 사용하지 마세요.",
        "danger",
    )
    add_table(
        doc,
        ["표시 코드/증상", "확인 순서"],
        [
            ("AGENT_DNS_FAILED", "입력한 Agent PC 이름 → 사내 DNS → IPv4 직접 입력"),
            ("AGENT_CONNECTION_REFUSED", "실제 Agent PC 주소 → 서비스 Running → 원격 TCP/18443"),
            ("AGENT_CLIENT_NOT_ALLOWED", "Viewer 주소가 10/8, 172.16/12 또는 192.168/16인지 확인"),
            ("AGENT_UNREACHABLE", "Viewer와 Agent PC 사이 라우팅 → 방화벽 → EDR 차단"),
            ("Agent/Viewer 버전 차이 경고", "API v4 호환이면 그대로 연결 가능. 문제 재현 시 같은 Release로 맞춤"),
            ("AGENT_PROTOCOL_MISMATCH", "Agent API가 v4인지 확인하고 같은 Release로 업데이트"),
            ("SETUP_PACKAGE_NOT_FOUND", "Agent ZIP 전체 압축 해제 → Setup과 Agent EXE·BUILD-MANIFEST 존재 확인"),
            ("SETUP_PACKAGE_HASH_MISMATCH", "실행 중지 → 공식 ZIP을 새 폴더에 다시 압축 해제 → EDR 격리 기록"),
            ("SETUP_SERVICE_FAILED", "Windows 서비스 관리 권한 → 기존 SamsungSwitchWatchAgent 상태"),
            ("SETUP_PATH_NOT_WRITABLE", "제품 폴더 권한·파일 상태 확인 실패 → 잠시 후 한 번만 재시도 → 반복되면 지원 코드 전달"),
            ("SETUP_PATH_INVALID / SETUP_PATH_UNTRUSTED", "반복 설치와 폴더·ACL 수동 변경 중지 → 지원 코드 전달"),
            ("FIREWALL_OVERLAP_PROTECTED", "외부 규칙은 보존됨 → 설치 계속 → Viewer 연결 테스트"),
            ("FIREWALL_REMOTE_ACCESS_UNCONFIRMED", "Agent 설치 유지 확인 → Viewer 연결 테스트 → TCP 실패 시 방화벽 서비스·Domain/Private·회사 GPO 확인"),
            ("AGENT_LOCAL_CONNECTION_UNCONFIRMED", "설치는 유지됨 → Viewer 연결 테스트 → 실패 시 Agent 서비스·로컬 HTTPS·EDR 확인"),
            ("SETUP_RECOVERY_REQUIRED", "설치/업데이트를 누르지 말고 '이전 상태 복구' 실행 → 복구 완료 확인 → 설치/업데이트를 별도로 시작"),
            ("SETUP_ROLLBACK_FAILED", "'진단정보 복사'로 최초 실패 코드와 ROLLBACK_* 단계 확인 → 증거 폴더를 그대로 보존 → Windows 관리자에게 전달"),
            ("ROLLBACK_STATE_MISMATCH", "복구 기록과 현재 설치 상태가 다름. 복구·설치 재시도와 파일 수동 정리를 중지"),
            ("ROLLBACK_SERVICE_STOP_FAILED", "SamsungSwitchWatchAgent 중지 상태와 서비스 제어 권한 확인"),
            ("ROLLBACK_FILE_RESTORE_FAILED / ROLLBACK_DATA_CLEANUP_FAILED", "Program Files·ProgramData 권한, 파일 잠금, 백신·EDR 격리 확인. 증거 폴더는 보존"),
            ("ROLLBACK_SERVICE_RESTORE_FAILED", "파일 복구 완료 여부를 먼저 확인한 뒤 기존 서비스 설정·시작 실패를 Windows 관리자에게 전달"),
            ("ROLLBACK_HTTPS_FIREWALL_RESTORE_FAILED / ROLLBACK_LEGACY_FIREWALL_RESTORE_FAILED", "제품 HTTPS 규칙과 이전 legacy 규칙의 개별 복원 실패. 규칙 범위를 수동 확대하지 않음"),
            ("ROLLBACK_JOURNAL_WRITE_FAILED / ROLLBACK_EVIDENCE_CLEANUP_FAILED", "journal 또는 복구 증거 정리 단계 실패. 아래 대상별 안전 코드를 확인하고 남은 기록·폴더는 보존"),
            ("ROLLBACK_STAGING_CLEANUP_FAILED / ROLLBACK_BACKUP_CLEANUP_FAILED", "현재 작업의 staging 또는 backup 자료 정리 실패. 수동 삭제 없이 익명 진단 전달"),
            ("ROLLBACK_FAILED_DIRECTORY_CLEANUP_FAILED / ROLLBACK_JOURNAL_CLEANUP_FAILED", "현재 작업의 failed 자료 또는 journal 정리·삭제 확인 실패. 설치는 잠긴 상태로 유지"),
            ("TARGET_NOT_ALLOWED", "장비 IPv4가 10/8, 172.16/12 또는 192.168/16인지 확인"),
            ("TCP_TIMEOUT", "Agent PC에서 장비 TCP/23 경로, ACL, 장비 Telnet 상태 확인"),
            ("AUTH_FAILED", "감시를 즉시 차단함. ID/PW와 login local 적용 여부 확인"),
            ("ENABLE_FAILED", "enable 필요 여부와 enable PW, 로그인 직후 프롬프트 확인"),
            ("QUERY_COMMAND_BLOCKED", "한 줄 show 형식, 줄바꿈/구분자 포함 여부 확인"),
            ("COMMAND_TIMEOUT", "30초 무응답 또는 90초 전체 제한. 실패한 수집 항목과 안전한 단계 확인"),
            ("TELNET_SESSION_CLOSED", "재연결 1회 뒤에도 종료됨. 완료된 결과와 남은 명령을 확인"),
            ("PROMPT_PARSE_FAILED", "로그인 후 > 또는 # 프롬프트 형식을 확인"),
            ("OUTPUT_LIMIT_EXCEEDED", "세션 처리 안전 한도 초과. 더 좁은 show 명령 사용"),
            ("VIEWER_DEVICE_STORE_CORRUPT", "원본은 .corrupt-*로 격리되고 이번 실행의 감시·명령이 중지됨. 격리 파일 보존 → Viewer 재시작 → 장비 재등록"),
            ("VIEWER_DEVICE_STORE_VERSION_UNSUPPORTED", "원본을 변경하지 말고 같거나 최신 Viewer를 설치해 다시 시작"),
            ("VIEWER_DEVICE_STORE_UNAVAILABLE", "감시·명령 중지됨. 사용자 폴더 권한·잠금·Windows 사용자 프로필 확인 후 Viewer 재시작"),
            ("VIEWER_DEVICE_STORE_WRITE_FAILED", "입력·원본은 유지됨. 사용자 폴더 권한·잠금·디스크 확인 후 Viewer 재시작"),
            ("VIEWER_SETTINGS_WRITE_FAILED", "Agent 연결은 유지됨. Viewer 사용자 폴더 권한과 디스크 여유 공간 확인"),
            ("VIEWER_MONITOR_STATE_CORRUPT", "이번 실행의 자동 감시가 중지됨. 격리된 원본을 보존하고 Viewer를 다시 시작해 새 기준선 생성"),
            ("VIEWER_MONITOR_STATE_VERSION_UNSUPPORTED", "파일을 변경하지 말고 같거나 최신 버전의 Viewer로 실행"),
            ("VIEWER_MONITOR_STATE_UNAVAILABLE", "자동 감시가 중지됨. Viewer 사용자 폴더 권한·파일 잠금 확인 후 재시작"),
            ("VIEWER_MONITOR_STATE_WRITE_FAILED", "장비 설정은 저장될 수 있음. 중복 등록하지 말고 감시 이력 파일 권한·잠금·디스크 확인"),
            ("VIEWER_MONITOR_CYCLE_FAILED", "다음 주기 재시도를 기다리고 반복되면 Viewer 진단 로그 확인"),
        ],
        [4300, 5060],
        header_size=8.5,
        body_size=8.25,
        body_line=1.0,
    )
    add_callout(
        doc,
        "Agent 폴더 신뢰 오류",
        "SETUP_PATH_NOT_WRITABLE은 기존 Agent 제품 폴더의 권한 또는 파일 상태를 확인하지 "
        "못했다는 뜻이며 로컬 HTTPS나 방화벽 오류가 아닙니다. 잠시 후 한 번만 재시도하고 "
        "반복되면 지원 코드만 전달하세요. SETUP_PATH_INVALID는 경로 형식 오류, "
        "SETUP_PATH_UNTRUSTED는 소유권 또는 junction·symlink 신뢰 검사 실패입니다. "
        "어떤 경우에도 제품 폴더와 ACL을 수동으로 삭제·변경하지 마세요.",
        "danger",
    )
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    add_callout(
        doc,
        "이전 상태 복구가 필요한 경우",
        "Setup은 남아 있는 journal을 읽기 전용으로 확인하고 SETUP_RECOVERY_REQUIRED를 표시합니다. "
        "이때 설치/업데이트는 잠겨 있으며 복구가 자동으로 시작되지 않습니다. '이전 상태 복구'를 "
        "직접 실행하세요. 정리 대상 삭제와 새 journal 검사가 모두 성공한 경우에만 설치 버튼이 "
        "활성화되며, 결과를 확인한 뒤 설치/업데이트를 별도로 시작합니다. 실패하면 "
        "'진단정보 복사'로 최초 실패 코드와 실제 경로가 없는 ROLLBACK_* 안전 단계를 전달합니다. "
        ".__staging_, .__backup_, .__failed_ 폴더와 journal은 원인 확인 증거이므로 "
        "삭제·이동·이름 변경하지 마세요.",
        "danger",
    )
    add_callout(
        doc,
        "방화벽 검증 실패 시",
        "FIREWALL_REMOTE_ADDRESS_MISMATCH 같은 코드는 어떤 필드가 안전 기준과 달랐는지만 "
        "알려 주며 실제 주소나 규칙 원문은 포함하지 않습니다. "
        "FIREWALL_REMOTE_ACCESS_UNCONFIRMED이면 Agent 설치는 유지되며 Viewer 연결 테스트로 "
        "원격 TCP/18443을 확인합니다. 실패하면 코드만 Windows 관리자에게 전달하세요.",
        "warning",
    )
    add_heading(doc, "현장 진단", 2, heading_num_id, page_break_before=True)
    add_body(
        doc,
        "Agent PC에서는 Windows 서비스에서 Agent 상태를 확인하고 필요하면 Agent Setup을 다시 실행합니다. Viewer에서는 Agent 연결 "
        "진단을 실행합니다. 두 화면의 단계와 안정 오류 코드만 기록하고 실제 IP, 계정, 비밀번호와 "
        "장비 원문은 외부로 전달하지 않습니다.",
    )
    add_table(
        doc,
        ["확인 위치", "정상 기준/조치"],
        [
            (
                "Viewer",
                "Agent 연결에 Agent PC 주소를 입력합니다. 원격 구성에서는 스위치 IP나 Viewer PC 주소를 사용하지 않습니다.",
            ),
            (
                "Agent PC",
                "Agent Setup의 설치/업데이트 결과에서 패키지, 서비스, 방화벽과 Agent 준비 상태를 확인합니다.",
            ),
            (
                "Viewer PC",
                "연결 설정의 주소 → DNS/IPv4 → TCP/18443 → HTTPS → Agent API/버전 단계 중 처음 실패한 항목을 확인합니다.",
            ),
            (
                "동일 PC 시험",
            "Agent 주소에 localhost 또는 해당 PC의 사설 IPv4를 입력해 Agent/API까지만 확인합니다. 스위치와 원격 Viewer 경로는 별도 확인합니다.",
            ),
        ],
        [1900, 7460],
        header_size=8.5,
        body_size=8.25,
        body_line=1.0,
    )
    add_callout(
        doc,
        "공유 전 확인",
        "익명 진단 TXT에는 원문이 없어야 합니다. 실제 회사 IP, 계정명, 비밀번호, 장비 출력이 보이면 공유하지 마세요.",
        "warning",
    )
    add_heading(
        doc,
        "업데이트·종료·운영 체크",
        1,
        heading_num_id,
        page_break_before=False,
    )
    add_heading(doc, "오프라인 수동 업데이트", 2, heading_num_id)
    add_steps(
        doc,
        [
            "새 Agent/Viewer ZIP을 승인된 경로로 전달하고 각각 임시 폴더에 압축 해제합니다.",
            "Agent PC에서 SamsungSwitchWatch.Agent.Setup.exe를 실행합니다. 입력 없이 설치/업데이트를 진행합니다.",
            "Viewer PC에서 새 Viewer 패키지의 SamsungSwitchWatch.Viewer.exe를 직접 실행합니다.",
            "Agent 연결, 장비 목록, 로그인 확인, 수집 진단, show 명령과 주기 감시를 순서대로 확인합니다.",
        ],
    )
    add_heading(doc, "Agent 제거", 2, heading_num_id)
    add_body(
        doc,
        "서비스와 방화벽을 변경하므로 Agent 제거는 사내 Windows 관리자가 승인된 관리 절차로 "
        "수행합니다. 공개 패키지에는 제거 스크립트를 포함하지 않습니다. 설치 폴더와 서비스는 "
        "별도 승인 없이 수동 삭제하지 않습니다.",
    )
    add_callout(
        doc,
        "설치 자료 삭제",
        "Program Files의 Agent 폴더와 Windows 서비스를 수동 삭제하면 업데이트와 복구가 어려워질 수 있습니다.",
        "danger",
    )
    add_heading(
        doc,
        "최종 운영 체크리스트",
        2,
        heading_num_id,
        page_break_before=False,
    )
    add_bullets(
        doc,
        [
            "Agent 서비스가 창 없이 Running이고 Viewer가 실제 사설 IPv4로 HTTPS 연결되는지 확인",
            "동일 PC 연결 확인 뒤 실제 원격 Viewer에서 Agent PC 주소로 다시 연결 확인",
            "계정은 Viewer에만 저장되고 로그인 확인 실패 장비의 주기 감시는 꺼지는지 확인",
            "show port status와 syslog 명령이 모델별로 동작하는지 확인",
            "민감 출력 비저장·Viewer 재실행·장애/복구·중복 억제는 모의 검증 후 사내 장비로 확인",
        ],
        bullet_num_id,
        font_size=10,
        space_after=2,
        line=1.05,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--images", type=Path, required=True)
    args = parser.parse_args()
    build_manual(args.output.resolve(), args.images.resolve())
    print(args.output.resolve())


if __name__ == "__main__":
    main()
